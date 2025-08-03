from langchain.document_loaders import PyPDFLoader,WebBaseLoader,TextLoader
from langchain.schema import Document
import docx   # ✅ Added for Word file support

def load_user_documents(sources):
    documents=[]

    for src in sources:
        try:
            if src.lower().endswith('.pdf'):
                load_pdf=PyPDFLoader(src)
                documents.extend(load_pdf.load())

            elif src.lower().startswith('http'):
                load_web_url=WebBaseLoader(src)
                documents.extend(load_web_url.load())
            
            elif src.lower().endswith('.txt'):
                load_text=TextLoader(src)
                documents.extend(load_text.load())

            elif src.lower().endswith('.docx'):  # ✅ Added Word file support
                word_doc = docx.Document(src)
                full_text = "\n".join([p.text for p in word_doc.paragraphs])
                documents.append(Document(page_content=full_text, metadata={"source": src}))

            else:
                # ✅ Wrap raw strings into Document to avoid AttributeError
                documents.append(Document(page_content=str(src)))

        except Exception as e:
            raise ValueError(f'⚠ Documents either not uploadedor not valid !!! \n'
                             f' Please Upload valid documents to continue chatting: {e}')
        
    return documents
